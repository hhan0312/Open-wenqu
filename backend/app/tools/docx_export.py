from __future__ import annotations

import io
from typing import Any

from docx import Document as DocxDocument
from docx.shared import Pt

from app.models.domain import GeneratedQuestion


def _questions_from_payload(payload: dict[str, Any]) -> list[GeneratedQuestion]:
    questions_raw = payload.get("questions") or []
    return [GeneratedQuestion.model_validate(q) for q in questions_raw]


def build_student_docx(*, title: str, passage: str, questions: list[GeneratedQuestion]) -> bytes:
    doc = DocxDocument()
    doc.add_heading(title, level=1)
    doc.add_heading("学生版", level=2)
    doc.add_paragraph("阅读材料")
    for para in passage.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    doc.add_paragraph("")
    doc.add_paragraph("题目")
    for i, q in enumerate(questions, start=1):
        doc.add_paragraph(f"{i}. {q.stem}")
        for opt in sorted(q.options, key=lambda x: x.key):
            doc.add_paragraph(f"  {opt.key}. {opt.text}")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_teacher_docx(*, title: str, passage: str, questions: list[GeneratedQuestion]) -> bytes:
    doc = DocxDocument()
    doc.add_heading(title, level=1)
    doc.add_heading("教师版", level=2)
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    doc.add_paragraph("阅读材料")
    for para in passage.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    doc.add_paragraph("")
    doc.add_paragraph("题目")
    for i, q in enumerate(questions, start=1):
        doc.add_paragraph(f"{i}. {q.stem}")
        for opt in sorted(q.options, key=lambda x: x.key):
            doc.add_paragraph(f"  {opt.key}. {opt.text}")
        doc.add_paragraph(f"正确答案：{q.correct_answer}")
        doc.add_paragraph("中文解析：")
        doc.add_paragraph(q.explanation_zh)
        doc.add_paragraph("原文证据：")
        doc.add_paragraph(q.evidence_text)
        if q.evidence_span:
            sp = q.evidence_span
            doc.add_paragraph(
                f"证据定位：block={sp.block_id} offsets={sp.start_offset}-{sp.end_offset} confidence={sp.confidence}"
            )
        doc.add_paragraph("干扰项分析：")
        for d in q.distractor_reviews:
            line = f"- {d.option_key}: {d.why_wrong_zh}"
            if d.confusion_risk_zh:
                line += f"（风险：{d.confusion_risk_zh}）"
            doc.add_paragraph(line)
        doc.add_paragraph("质量自评：")
        doc.add_paragraph(
            f"评分 {q.quality.score}；清晰性：{q.quality.clarity_zh}；难度匹配：{q.quality.difficulty_match_zh}；答案唯一性：{q.quality.uniqueness_zh}"
        )
        if q.quality.issues_zh:
            doc.add_paragraph("问题：" + "；".join(q.quality.issues_zh))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_student_teacher_docx(*, title: str, passage: str, payload: dict[str, Any]) -> bytes:
    """Single downloadable document containing student section then teacher section."""
    questions = _questions_from_payload(payload)
    doc = DocxDocument()
    doc.add_heading(title, level=1)

    doc.add_heading("学生版", level=2)
    doc.add_paragraph("阅读材料")
    for para in passage.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    doc.add_paragraph("")
    doc.add_paragraph("题目")
    for i, q in enumerate(questions, start=1):
        doc.add_paragraph(f"{i}. {q.stem}")
        for opt in sorted(q.options, key=lambda x: x.key):
            doc.add_paragraph(f"  {opt.key}. {opt.text}")

    doc.add_page_break()
    doc.add_heading("教师版", level=2)
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    doc.add_paragraph("阅读材料")
    for para in passage.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    doc.add_paragraph("")
    doc.add_paragraph("题目与解析")
    for i, q in enumerate(questions, start=1):
        doc.add_paragraph(f"{i}. {q.stem}")
        for opt in sorted(q.options, key=lambda x: x.key):
            doc.add_paragraph(f"  {opt.key}. {opt.text}")
        doc.add_paragraph(f"正确答案：{q.correct_answer}")
        doc.add_paragraph("中文解析：")
        doc.add_paragraph(q.explanation_zh)
        doc.add_paragraph("原文证据：")
        doc.add_paragraph(q.evidence_text)
        if q.evidence_span:
            sp = q.evidence_span
            doc.add_paragraph(
                f"证据定位：block={sp.block_id} offsets={sp.start_offset}-{sp.end_offset} confidence={sp.confidence}"
            )
        doc.add_paragraph("干扰项分析：")
        for d in q.distractor_reviews:
            line = f"- {d.option_key}: {d.why_wrong_zh}"
            if d.confusion_risk_zh:
                line += f"（风险：{d.confusion_risk_zh}）"
            doc.add_paragraph(line)
        doc.add_paragraph("质量自评：")
        doc.add_paragraph(
            f"评分 {q.quality.score}；清晰性：{q.quality.clarity_zh}；难度匹配：{q.quality.difficulty_match_zh}；答案唯一性：{q.quality.uniqueness_zh}"
        )
        if q.quality.issues_zh:
            doc.add_paragraph("问题：" + "；".join(q.quality.issues_zh))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
